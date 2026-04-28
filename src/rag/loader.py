from docx import Document as DocxDocument
from langchain_core.documents import Document
from src.utils.reindexing import get_file_hash, load_hashes
from pathlib import Path

hashes = load_hashes()

def upload_and_load_docs(file) -> Document:

    print(f"Processing file: {file}")
    file_hash = get_file_hash(Path(file))
    hash_on_disk = hashes.get(file, None)
    if hash_on_disk and hash_on_disk == file_hash:
        print(f"File {file} has not changed. Skipping processing.")
        return None
    else:
        if hash_on_disk:
            print(f"File {file} has changed. Updating hash and reprocessing.")

    
    doc = DocxDocument(file)
    metadata = {"file_name": file}

    if doc.tables:
        first_table = doc.tables[0]
        for row in first_table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) == 2:
                key,value = cells
                metadata[key] = value
            elif len(cells) > 2:
                metadata[cells[0]] = " | ".join(cells[1:])
    else:
        print("No tables found in the document!")

    metadata.pop('Metadata',0)

    content_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            content_parts.append(para.text.strip())

    for _, table in enumerate(doc.tables[1:],start=1):
        table_data = []
        for row in table.rows:
            table_data.append([cell.text.strip() for cell in row.cells])

        table_text = "\n".join([" | ".join(row) for row in table_data])
        content_parts.append(table_text)

    full_content = "\n\n".join(content_parts)

    if metadata.get('Applicable Region') in ["", None, " "]:
        metadata['Applicable Region'] = "Global"

    metadata["Applicable Region"] = metadata["Applicable Region"].replace(" ","").lower()

    return Document(
        page_content = full_content,
        metadata=metadata
    )
